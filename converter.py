#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_page_numbers(doc):
    """Add page numbers to the footer"""
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""
    
    # Add page number
    run = footer_para.add_run()
    
    # Create page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    
    # Format footer paragraph
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.paragraph_format.space_before = Pt(12)
    
    # Style the page number
    for run in footer_para.runs:
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(128, 128, 128)

def add_page_break(doc):
    """Add a page break to the document"""
    doc.add_page_break()

def set_document_margins(doc):
    """Set professional document margins"""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def create_cover_page(doc):
    """Create a professional cover page"""
    
    # Add spacing at the top
    for _ in range(6):
        doc.add_paragraph()
    
    # Main title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("O QUE CREMOS")
    run.font.name = 'Garamond'
    run.font.size = Pt(44)
    run.font.bold = True
    run.font.color.rgb = RGBColor(44, 62, 80)
    title.paragraph_format.space_after = Pt(6)
    
    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Fundamentos Doutrinários da Fé Cristã")
    run.font.name = 'Garamond'
    run.font.size = Pt(18)
    run.font.italic = True
    run.font.color.rgb = RGBColor(74, 90, 108)
    subtitle.paragraph_format.space_after = Pt(24)
    
    # Separator line
    sep1 = doc.add_paragraph()
    sep1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sep1.add_run("━" * 40)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(200, 200, 200)
    sep1.paragraph_format.space_after = Pt(24)
    
    # Spacing
    for _ in range(5):
        doc.add_paragraph()
    
    # Church info
    church = doc.add_paragraph()
    church.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = church.add_run("Igreja Batista Pedras Vivas")
    run.font.name = 'Calibri'
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(52, 73, 94)
    church.paragraph_format.space_after = Pt(12)
    
    # Course info
    course = doc.add_paragraph()
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course.add_run("Apostila de Estudos - Escola Bíblica Dinâmica")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(74, 90, 108)
    course.paragraph_format.space_after = Pt(6)
    
    # Description
    description = doc.add_paragraph()
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = description.add_run("18 Aulas sobre as Doutrinas Fundamentais da Fé Cristã")
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(100, 110, 120)
    description.paragraph_format.space_after = Pt(24)
    
    # Spacing
    for _ in range(6):
        doc.add_paragraph()
    
    # Year
    year = doc.add_paragraph()
    year.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = year.add_run("2026")
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(100, 110, 120)
    year.paragraph_format.space_after = Pt(12)
    
    # Page break after cover
    doc.add_page_break()

def add_fill_line(doc):
    """Add a fillable line for student responses"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    
    # Add a dotted/underlined fill line
    run = p.add_run("_" * 85)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(180, 180, 180)

def parse_markdown_to_docx(markdown_file, docx_output):
    """Convert markdown file to DOCX with professional formatting"""
    
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set professional margins
    set_document_margins(doc)
    
    # Create title page
    create_cover_page(doc)
    
    lines = content.split('\n')
    
    i = 0
    last_was_empty = False
    
    while i < len(lines):
        line = lines[i]
        
        # Check for page break before AULA X (but not AULA 1)
        if re.match(r'^# AULA \d+', line) and i > 0:
            match = re.match(r'^# AULA (\d+)', line)
            if match and int(match.group(1)) > 1:
                add_page_break(doc)
        
        # Check for fill lines (underscores)
        if re.match(r'^_{4,}', line):
            add_fill_line(doc)
            last_was_empty = False
            i += 1
            continue
        
        # Process different heading levels
        if re.match(r'^# AULA', line):  # Main AULA heading
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            p.paragraph_format.line_spacing = 1.15
            
            # Style the heading
            for run in p.runs:
                run.font.name = 'Garamond'
                run.font.size = Pt(26)
                run.font.bold = True
                run.font.color.rgb = RGBColor(44, 62, 80)
            
            last_was_empty = False
        
        elif re.match(r'^# [A-Z]', line):  # Other main headings
            title = line[2:].strip()
            p = doc.add_heading(title, level=1)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(12)
            
            for run in p.runs:
                run.font.name = 'Garamond'
                run.font.size = Pt(22)
                run.font.bold = True
                run.font.color.rgb = RGBColor(44, 62, 80)
            
            last_was_empty = False
        
        elif re.match(r'^## ', line):  # Section headings
            title = line[3:].strip()
            p = doc.add_heading(title, level=2)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.line_spacing = 1.15
            
            for run in p.runs:
                run.font.name = 'Garamond'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(52, 73, 94)
            
            last_was_empty = False
        
        elif re.match(r'^### ', line):  # Subsection headings
            title = line[4:].strip()
            p = doc.add_heading(title, level=3)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            
            for run in p.runs:
                run.font.name = 'Garamond'
                run.font.size = Pt(13)
                run.font.bold = True
                run.font.color.rgb = RGBColor(74, 90, 108)
            
            last_was_empty = False
        
        elif re.match(r'^#### ', line):  # Level 4 headings
            title = line[5:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.left_indent = Inches(0)
            p.paragraph_format.line_spacing = 1.15
            
            run = p.add_run(title)
            run.font.name = 'Garamond'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(100, 110, 120)
            
            last_was_empty = False
        
        elif line == '---':  # Horizontal rule / divider
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr')
            if pBdr is None:
                from docx.oxml import parse_xml
                pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="1" w:color="D0D0D0"/></w:pBdr>')
                pPr.append(pBdr)
            
            last_was_empty = False
        
        elif re.match(r'^- ', line):  # Bullet points
            text = line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.line_spacing = 1.15
            
            format_inline_markdown(p, text)
            
            # Style the bullet text
            for run in p.runs:
                if not run.bold and not run.italic:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(51, 51, 51)
            
            last_was_empty = False
        
        elif re.match(r'^\d+\. ', line):  # Numbered lists
            match = re.match(r'^(\d+)\. (.*)$', line)
            if match:
                text = match.group(2).strip()
                p = doc.add_paragraph(style='List Number')
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(4)
                p.paragraph_format.line_spacing = 1.15
                
                format_inline_markdown(p, text)
                
                for run in p.runs:
                    if not run.bold and not run.italic:
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(51, 51, 51)
                
                last_was_empty = False
        
        elif line.strip() == '':  # Empty lines - skip multiple consecutive ones
            if not last_was_empty:
                # Add a single empty paragraph but without adding spacing
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                last_was_empty = True
        
        else:  # Regular paragraphs
            if line.strip():
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(8)
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                format_inline_markdown(p, line.strip())
                
                # Apply default text formatting
                for run in p.runs:
                    if not run.bold and not run.italic:  # Don't override formatted text
                        run.font.name = 'Calibri'
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(51, 51, 51)
                
                last_was_empty = False
        
        i += 1
    
    # Add page numbers to all sections
    add_page_numbers(doc)
    
    # Save document
    doc.save(docx_output)
    print(f"[OK] Documento profissional criado: {docx_output}")

def format_inline_markdown(paragraph, text):
    """Add text to paragraph with inline markdown formatting"""
    
    # Pattern to split by markdown formatting
    pattern = r'(\*\*.*?\*\*|\*.*?\*|__.*?__|_.*?_)'
    
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        if part.startswith('**') and part.endswith('**'):
            # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(44, 62, 80)
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(51, 51, 51)
        elif part.startswith('__') and part.endswith('__'):
            # Bold (alt syntax)
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(44, 62, 80)
        elif part.startswith('_') and part.endswith('_') and len(part) > 2:
            # Italic (alt syntax)
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(51, 51, 51)
        else:
            # Regular text
            run = paragraph.add_run(part)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(51, 51, 51)

if __name__ == '__main__':
    markdown_file = r'c:\Users\ederf\OneDrive\Eder\ibpv\EBD-2026\docs\EBD-2026-O-Que-Cremos.md'
    docx_output = r'c:\Users\ederf\OneDrive\Eder\ibpv\EBD-2026\docs\EBD-2026-O-Que-Cremos.docx'
    
    parse_markdown_to_docx(markdown_file, docx_output)
